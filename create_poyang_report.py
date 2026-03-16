"""
创建鄱阳工单运维报告
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook

# 读取鄱阳工单数据
print("读取鄱阳工单数据...")
wb = load_workbook(r"D:\XLX\XLXWJ\鄱阳工单数据_1773391976677.xlsx")
ws = wb.active

# 获取所有数据
data = []
headers = []
for i, row in enumerate(ws.iter_rows(values_only=True), 1):
    if i == 1:
        headers = row
    else:
        data.append(row)

print(f"共 {len(data)} 条工单")

# 创建Word文档
doc = Document()

# 添加标题
title = doc.add_paragraph()
title.add_run("鄱阳县中医院云净血透系统运维报告")
title.runs[0].font.size = Pt(18)
title.runs[0].font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加表格
table = doc.add_table(rows=len(data)+1, cols=len(headers))
table.style = 'Table Grid'

# 添加表头
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = str(header)
    cell.paragraphs[0].runs[0].font.bold = True

# 添加数据
for i, row in enumerate(data):
    for j, cell_data in enumerate(row):
        cell = table.rows[i+1].cells[j]
        if cell_data is not None:
            cell.text = str(cell_data)

# 保存文档
output_file = r"D:\XLX\XLXWJ\鄱阳中医院云净血透系统运维报告.docx"
doc.save(output_file)

print(f"\n报告已保存: {output_file}")
